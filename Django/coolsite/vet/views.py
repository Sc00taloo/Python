from django.contrib.auth import get_user_model, logout, login
from django.contrib.auth.decorators import login_required
#from django.contrib.auth.forms import AuthenticationForm
from django.contrib.auth.views import LoginView
from django.http import HttpResponse
from django.shortcuts import render, redirect, get_object_or_404
#from django.http import HttpResponseRedirect
from django.urls import reverse_lazy
from django.views.generic import CreateView

import docx
from openpyxl import Workbook
import aspose.pdf as ap

from .models import *
from .form import *


def index(request):
    return render(request, 'vet/index.html')
# def login(request):
#     return render(request, 'vet/login.html')

def logout_user(request):
    logout(request)
    return redirect('login')
def hosts(request):
    if request.method == 'POST':
        return redirect('http://127.0.0.1:8000/')
    return render(request, 'vet/host.html')
def hword(request):
    doc = docx.Document()
    data = list_of_hosts.objects.all().values()
    str_data = [str(row) for row in data]
    for list in str_data:
    # добавляем первый параграф
        doc.add_paragraph(list)
    doc.save('hosts.docx')
    return HttpResponse('Тест')
def hexcel(request):
    data = list_of_hosts.objects.all()
    wb = Workbook()
    sheet = wb.active
    sheet['A1'] = 'Имя'
    sheet['B1'] = 'Фамилия'
    sheet['C1'] = 'Адресс'
    sheet['D1'] = 'Телефон'
    for row, item in enumerate(data, start=2):
        sheet.cell(row=row, column=1, value=item.first_name)
        sheet.cell(row=row, column=2, value=item.last_name)
        sheet.cell(row=row, column=3, value=item.address)
        sheet.cell(row=row, column=4, value=item.phone)
    wb.save('host.xlsx')
    return HttpResponse('Тест')
def hpdf(request):
    data = list_of_hosts.objects.all().values()
    str_data = [str(row) for row in data]
    # Инициализировать объект документа
    document = ap.Document()
    # Добавить страницу
    page = document.pages.add()
    # Инициализировать объект textfragment
    for list in str_data:
        text_fragment = ap.text.TextFragment(str_data)
        # Добавить фрагмент текста на новую страницу
        page.paragraphs.add(text_fragment)
    document.save("host.pdf")
    return HttpResponse('Тест')

def animals(request):
    if request.method == 'POST':
        return redirect('http://127.0.0.1:8000/')
    return render(request, 'vet/animal.html')
def aword(request):
    doc = docx.Document()
    data = list_of_animals.objects.all().values()
    str_data = [str(row) for row in data]
    for list in str_data:
    # добавляем первый параграф
        doc.add_paragraph(list)
    doc.save('animal.docx')
    return HttpResponse('Тест')
def aexcel(request):
    data = list_of_animals.objects.all()
    wb = Workbook()
    sheet = wb.active
    sheet['A1'] = 'Имя'
    sheet['B1'] = 'Тип'
    sheet['C1'] = 'Порода'
    sheet['D1'] = 'Лет'
    sheet['E1'] = 'Хозяин'
    for row, item in enumerate(data, start=2):
        sheet.cell(row=row, column=1, value=item.name)
        sheet.cell(row=row, column=2, value=item.kind_of_animal)
        sheet.cell(row=row, column=3, value=item.breed)
        sheet.cell(row=row, column=4, value=item.age)
        sheet.cell(row=row, column=5, value=item.host_id)
    wb.save('animal.xlsx')
    return HttpResponse('Тест')
def apdf(request):
    data = list_of_animals.objects.all().values()
    str_data = [str(row) for row in data]
    # Инициализировать объект документа
    document = ap.Document()
    # Добавить страницу
    page = document.pages.add()
    # Инициализировать объект textfragment
    for list in str_data:
        text_fragment = ap.text.TextFragment(str_data)
        # Добавить фрагмент текста на новую страницу
        page.paragraphs.add(text_fragment)
    document.save("animal.pdf")
    return HttpResponse('Тест')

def diseasis(request):
    if request.method == 'POST':
        return redirect('http://127.0.0.1:8000/')
    return render(request, 'vet/diseases.html')
def dword(request):
    doc = docx.Document()
    data = diseases.objects.all().values()
    str_data = [str(row) for row in data]
    for list in str_data:
    # добавляем первый параграф
        doc.add_paragraph(list)
    doc.save('diseases.docx')
    return HttpResponse('Тест')
def dexcel(request):
    data = diseases.objects.all()
    wb = Workbook()
    sheet = wb.active
    sheet['A1'] = 'Имя'
    for row, item in enumerate(data, start=2):
        sheet.cell(row=row, column=1, value=item.name)
    wb.save('diseases.xlsx')
    return HttpResponse('Тест')
def dpdf(request):
    data = diseases.objects.all().values()
    str_data = [str(row) for row in data]
    # Инициализировать объект документа
    document = ap.Document()
    # Добавить страницу
    page = document.pages.add()
    # Инициализировать объект textfragment
    for list in str_data:
        text_fragment = ap.text.TextFragment(str_data)
        # Добавить фрагмент текста на новую страницу
        page.paragraphs.add(text_fragment)
    document.save("diseases.pdf")
    return HttpResponse('Тест')

@login_required
def hosts_select(request):
    if request.method == 'POST':
        return redirect('http://127.0.0.1:8000/hosts')
    data = list_of_hosts.objects.filter(user_id=request.user.id)
    if request.user.is_superuser:
        return render(request, 'vet/hselect.html', {'data':list_of_hosts.objects.all()})
    elif any(host.user == request.user for host in data):
        return render(request, 'vet/hselect.html', {'data': data})
    return HttpResponse('Нет записей =(')
@login_required
def hosts_update(request):
    if request.method == 'POST':
        form = hostsUpdate(request.POST)
        if form.is_valid():
            id_update = form.cleaned_data['id']
            obj = list_of_hosts.objects.get(id=id_update)
            if request.user.is_superuser or obj.user == request.user:
                instance = get_object_or_404(list_of_hosts, id=id_update)
                if form.cleaned_data['first_name']:
                    instance.first_name = form.cleaned_data['first_name']
                if form.cleaned_data['last_name']:
                    instance.last_name = form.cleaned_data['last_name']
                if form.cleaned_data['address']:
                    instance.address = form.cleaned_data['address']
                if form.cleaned_data['phone']:
                    instance.phone = form.cleaned_data['phone']
                instance.save()
                return redirect('hosts')
            else:
                return HttpResponse('У вас нет разрешения на обновления этой таблицы')
    else:
        form = hostsUpdate()
    return render(request, 'vet/hupdate.html', {'form':form})
@login_required
def hosts_insert(request):
    if request.method == 'POST':
        form = hostsForm(request.POST)
        form.instance.user_id=request.user.id
        if form.is_valid():
            form.save()
            return redirect('hosts')
    else:
        form = hostsForm()
    return render(request, 'vet/hinsert.html', {'form':form})
@login_required
def hosts_delete(request):
    if request.method == 'POST':
        form = hostsDelete(request.POST)
        if form.is_valid():
            id_to_delete = form.cleaned_data['id']
            obj = list_of_hosts.objects.get(id=id_to_delete)
            if request.user.is_superuser or obj.user == request.user:
                obj.delete()
                return redirect('hosts')
            else:
                return HttpResponse('У вас нет разрешения на удаление этой таблицы')
    else:
        form = hostsDelete()
    return render(request, 'vet/hdelete.html', {'form':form})

@login_required
def animals_select(request):
    if request.method == 'POST':
        return redirect('http://127.0.0.1:8000/animals')
    data = list_of_animals.objects.filter(user_id=request.user.id)
    if request.user.is_superuser:
        return render(request, 'vet/aselect.html', {'data': list_of_animals.objects.all()})
    elif any(animal.user == request.user for animal in data):
        return render(request, 'vet/aselect.html', {'data': data})
    return HttpResponse('Нет записей =(')
@login_required
def animals_update(request):
    if request.method == 'POST':
        form = animalsUpdate(request.POST)
        if form.is_valid():
            id_update = form.cleaned_data['id']
            obj = list_of_animals.objects.get(id=id_update)
            if request.user.is_superuser or obj.user == request.user:
                id = form.cleaned_data['id']
                instance = get_object_or_404(list_of_animals, id=id)
                if form.cleaned_data['kind_of_animal']:
                    instance.kind_of_animal = form.cleaned_data['kind_of_animal']
                if form.cleaned_data['breed']:
                    instance.breed = form.cleaned_data['breed']
                if form.cleaned_data['age']:
                    instance.age = form.cleaned_data['age']
                if form.cleaned_data['host']:
                    instance.host = form.cleaned_data['host']
                instance.save()
                return redirect('animals')
            else:
                return HttpResponse('У вас нет разрешения на обновления этой таблицы')
    else:
        form = animalsUpdate()
    return render(request, 'vet/aupdate.html', {'form':form})
@login_required
def animals_insert(request):
    if request.method == 'POST':
        form = animalsForm(request.POST)
        form.instance.user_id=request.iser.id
        if form.is_valid():
            form.save()
            return redirect('animals')
    else:
        form = animalsForm()
    return render(request, 'vet/ainsert.html', {'form': form})
@login_required
def animals_delete(request):
    if request.method == 'POST':
        form = animalsDelete(request.POST)
        if form.is_valid():
            id_to_delete = form.cleaned_data['id']
            obj = list_of_animals.objects.get(id=id_to_delete)
            if request.user.is_superuser or obj.user == request.user:
                obj.delete()
                return redirect('animals')
            else:
                return HttpResponse('У вас нет разрешения на удаление этой таблицы')
    else:
        form = animalsDelete()
    return render(request, 'vet/adelete.html', {'form': form})

@login_required
def diseases_select(request):
    if request.method == 'POST':
        return redirect('http://127.0.0.1:8000/diseasis')
    data = diseases.objects.filter(user_id=request.user.id)
    if request.user.is_superuser:
        return render(request, 'vet/dselect.html', {'data': diseases.objects.all()})
    elif any(disease.user == request.user for disease in data):
        return render(request, 'vet/dselect.html', {'data': data})
    return HttpResponse('Нет записей =(')
@login_required
def diseases_update(request):
    if request.method == 'POST':
        form = diseasesUpdate(request.POST)
        if form.is_valid():
            id_update = form.cleaned_data['id']
            obj = diseases.objects.get(id=id_update)
            if request.user.is_superuser or obj.user == request.user:
                instance = get_object_or_404(diseases, id=id)
                if form.cleaned_data['name']:
                    instance.name = form.cleaned_data['name']
                instance.save()
                return redirect('http://127.0.0.1:8000/diseasis')
            else:
                return HttpResponse('У вас нет разрешения на обновления этой таблицы')
    else:
        form = diseasesUpdate()
    return render(request, 'vet/dupdate.html', {'form':form})
@login_required
def diseases_insert(request):
    if request.method == 'POST':
        form = diseasesForm(request.POST)
        form.instance.user_id=request.user.id
        if form.is_valid():
            form.save()
            return redirect('http://127.0.0.1:8000/diseasis')
    else:
        form = diseasesForm()
    return render(request, 'vet/dinsert.html', {'form':form})
@login_required
def diseases_delete(request):
    if request.method == 'POST':
        form = diseasesDelete(request.POST)
        if form.is_valid():
            id_to_delete = form.cleaned_data['id']
            obj = diseases.objects.get(id=id_to_delete)
            if request.user.is_superuser or obj.user == request.user:
                obj.delete()
                return redirect('http://127.0.0.1:8000/diseasis')
            else:
                return HttpResponse('У вас нет разрешения на удаление этой таблицы')
    else:
        form = diseasesDelete()
    return render(request, 'vet/ddelete.html', {'form':form})

class RegistrationView(CreateView):
    model = get_user_model()
    form_class = RegistrationForm
    template_name = 'vet/register.html'
    success_url = reverse_lazy('login')

    def get_context_data(self, *, object_list=None, **kwargs):
        context = super().get_context_data(**kwargs)
        return dict(list(context.items()))

    def form_valid(self, form):
        user = form.save()
        login(self.request, user)
        return redirect('home')

class LoginUser(LoginView):
    form_class = LoginUserForm
    template_name = 'vet/login.html'

    def get_context_data(self, *, object_list=None, **kwargs):
        context = super().get_context_data(**kwargs)
        return dict(list(context.items()))

    def get_success_url(self):
        return reverse_lazy('home')